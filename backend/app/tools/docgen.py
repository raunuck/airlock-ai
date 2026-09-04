from docx import Document
from datetime import datetime

def write_approval_note(findings: str, output_dir: str = "outputs") -> str:
    doc = Document()
    doc.add_heading("Approval Note", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("Findings summary:")
    doc.add_paragraph(findings)
    doc.add_paragraph("Recommended action: ___________")

    import os
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f"approval_note_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(path)
    return path